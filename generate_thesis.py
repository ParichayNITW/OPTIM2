from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

# ── OMML helpers ──────────────────────────────────────────────────────────────
M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _m(tag, *children, **attrib):
    el = OxmlElement(f"m:{tag}")
    for k, v in attrib.items():
        el.set(qn(f"m:{k}") if ":" not in k else k, v)
    for c in children:
        if c is not None:
            el.append(c)
    return el

def _mrun(text, bold=False, italic=True):
    rPr = _m("rPr")
    sty = _m("sty")
    sty.set(qn("m:val"), "bi" if (bold and italic) else ("b" if bold else ("i" if italic else "p")))
    rPr.append(sty)
    t = _m("t")
    t.text = text
    r = _m("r", rPr, t)
    return r

def _mtext(text):          # upright / plain text in math
    rPr = _m("rPr"); sty = _m("sty"); sty.set(qn("m:val"), "p"); rPr.append(sty)
    t = _m("t"); t.text = text
    return _m("r", rPr, t)

def _frac(num_el, den_el):
    return _m("f", _m("num", num_el), _m("den", den_el))

def _sup(base, exp):
    return _m("sSup", _m("e", base), _m("sup", exp))

def _sub(base, sb):
    return _m("sSub", _m("e", base), _m("sub", sb))

def _subsup(base, sb, sup_el):
    return _m("sSubSup", _m("e", base), _m("sub", sb), _m("sup", sup_el))

def _sqrt(inner):
    return _m("rad", _m("radPr", _m("degHide", **{qn("m:val"):"1"})), _m("deg"), _m("e", inner))

def _row(*els):
    return _m("r", *els)          # just wrap – caller unpacks

def insert_eq(doc, *math_els, eq_num=None, space_before=8, space_after=8):
    """Insert a display-mode OMML equation, optionally with eq number."""
    # Use a 1-row, 2-col table so the number sits at the right margin
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(13)
    tbl.columns[1].width = Cm(3)
    cell_eq  = tbl.cell(0, 0)
    cell_num = tbl.cell(0, 1)
    for cell in (cell_eq, cell_num):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top","left","bottom","right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"),"none"); b.set(qn("w:sz"),"0"); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto")
            tcBorders.append(b)
        tcPr.append(tcBorders)

    p = cell_eq.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(space_before * 20))
    spacing.set(qn("w:after"),  str(space_after  * 20))
    pPr.append(spacing)

    oMath = _m("oMath")
    for el in math_els:
        oMath.append(el)
    p._p.append(oMath)

    pn = cell_num.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pn.paragraph_format.space_before = Pt(space_before)
    pn.paragraph_format.space_after  = Pt(space_after)
    if eq_num:
        run = pn.add_run(f"({eq_num})")
        run.font.size = Pt(11)
    return tbl

# ── Document helpers ───────────────────────────────────────────────────────────
def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.5)
    sec.right_margin  = Cm(2.5)
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21.0)

    styles = doc.styles
    # Normal
    n = styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(12)
    pf = n.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

    for lvl, sz, bold, space_b in [("Heading 1",14,True,18),
                                    ("Heading 2",13,True,12),
                                    ("Heading 3",12,True, 8)]:
        h = styles[lvl]
        h.font.name  = "Times New Roman"
        h.font.size  = Pt(sz)
        h.font.bold  = bold
        h.font.color.rgb = RGBColor(0x00,0x00,0x00)
        h.paragraph_format.space_before = Pt(space_b)
        h.paragraph_format.space_after  = Pt(6)
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        h.paragraph_format.keep_with_next    = True
    return doc

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    p.runs[0].bold = True
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(13); r.bold = True
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12); r.bold = True
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p

def eq_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        r.font.size = Pt(10); r.italic = True
    return p

def tbl_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        r.bold = True; r.font.size = Pt(11)
    return p

def styled_table(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=ncols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = t.rows[0]
    hrow.height = Cm(0.8)
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = "Times New Roman"
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"1F497D")
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            if ri % 2 == 0:
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"DCE6F1")
                tcPr.append(shd)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    return t

def page_break(doc):
    doc.add_page_break()

def add_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom); pPr.append(pBdr)
    return p

def centered(doc, text, size=12, bold=False, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(size); r.bold = bold
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Cm(1.0 + level*0.5)
    r = p.add_run(text)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    return p

